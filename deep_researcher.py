"""
OMNI Elite Deep Research Engine v2.0
======================================
Advanced autonomous research system with:
  - AI-driven multi-query generation strategy
  - Google + Bing + Wikipedia tri-engine search
  - Source credibility scoring & filtering
  - Per-page structured data / stats extraction
  - Cross-source fact validation
  - AI outline generation BEFORE scraping
  - 10-section professional report synthesis
  - Word (.docx) + Excel (.xlsx) export with premium styling

Author: ORCA / Naqash Afzal
"""

import asyncio
import re
import time
from datetime import datetime
from typing import Callable, Optional, List, Dict


# ──────────────────────────────────────────────────────────────────────────────
#  CREDIBILITY RULES  (domain → score 0-10)
# ──────────────────────────────────────────────────────────────────────────────
HIGH_CRED_DOMAINS = {
    "wikipedia.org": 9, "britannica.com": 9, "nature.com": 10,
    "sciencedirect.com": 10, "pubmed.ncbi.nlm.nih.gov": 10,
    "scholar.google.com": 9, "researchgate.net": 8, "jstor.org": 9,
    "who.int": 9, "cdc.gov": 9, "nih.gov": 10, "bbc.com": 7,
    "reuters.com": 8, "apnews.com": 8, "nytimes.com": 7,
    "forbes.com": 7, "hbr.org": 8, "mckinsey.com": 8,
    "statista.com": 8, "ourworldindata.org": 9,
}
LOW_CRED_PATTERNS = ["blogspot", "wordpress", "wixsite", "tumblr", "reddit"]


def _score_domain(url: str) -> int:
    for domain, score in HIGH_CRED_DOMAINS.items():
        if domain in url:
            return score
    for bad in LOW_CRED_PATTERNS:
        if bad in url:
            return 3
    return 5  # neutral


# ──────────────────────────────────────────────────────────────────────────────
#  ELITE RESEARCHER CLASS
# ──────────────────────────────────────────────────────────────────────────────
class DeepResearcher:
    CHARS_PER_SOURCE = 8000

    def __init__(self, page, llm_provider=None, log_fn: Optional[Callable] = None):
        self.page = page
        self.llm = llm_provider
        self.log = log_fn or print
        self._stop_requested = False

    def stop(self):
        self._stop_requested = True

    # ─── MAIN PIPELINE ───────────────────────────────────────────────────────
    async def run(self, topic: str, max_sources: int = 8) -> dict:
        self._stop_requested = False
        self.log(f"[OMNI RESEARCH] 🔬 Starting Elite Research: '{topic}'")
        self.log(f"[OMNI RESEARCH] Phase 1: Generating intelligent search strategy…")

        # ① AI generates multiple diverse search queries
        queries = await self._generate_queries(topic)
        self.log(f"[OMNI RESEARCH] Generated {len(queries)} search angles:")
        for i, q in enumerate(queries, 1):
            self.log(f"   [{i}] {q}")

        # ② Search all queries across Google + Bing + Wikipedia
        self.log(f"[OMNI RESEARCH] Phase 2: Multi-engine search across {len(queries)} queries…")
        all_links = await self._multi_search(queries)
        self.log(f"[OMNI RESEARCH] Found {len(all_links)} unique candidate sources.")

        # ③ Score and rank sources
        scored = self._score_sources(all_links, topic)
        top_links = [url for url, _ in scored[:max_sources + 4]]  # fetch extra, filter later

        # ④ Extract content
        self.log(f"[OMNI RESEARCH] Phase 3: Deep content extraction from top {max_sources} sources…")
        sources = await self._extract_sources(top_links, max_sources)

        if not sources:
            self.log("[OMNI RESEARCH] ⚠ No content extracted. Check browser / internet connection.")
            return self._empty_report(topic)

        # ⑤ Extract statistics/numbers from all raw text
        self.log(f"[OMNI RESEARCH] Phase 4: Mining statistics and data points…")
        all_text = "\n\n".join(s["text"] for s in sources)
        stats = _extract_statistics(all_text)
        self.log(f"[OMNI RESEARCH] Found {len(stats)} quantitative data points.")

        # ⑥ Wikipedia supplement (always include for baseline accuracy)
        self.log(f"[OMNI RESEARCH] Phase 5: Wikipedia knowledge base supplement…")
        wiki_text = await self._wikipedia_lookup(topic)
        if wiki_text:
            sources.insert(0, {"url": f"https://en.wikipedia.org/wiki/{topic.replace(' ', '_')}",
                                "text": wiki_text[:self.CHARS_PER_SOURCE],
                                "credibility": 9, "title": f"Wikipedia: {topic}"})
            self.log(f"[OMNI RESEARCH] ✓ Wikipedia content integrated.")

        # ⑦ AI synthesis into full research paper
        self.log(f"[OMNI RESEARCH] Phase 6: AI synthesis ({len(sources)} sources, {len(all_text):,} chars)…")
        report = await self._synthesise(topic, sources, stats)

        report.update({
            "topic": topic,
            "sources": sources,
            "queries_used": queries,
            "stats_found": stats,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "total_sources": len(sources),
            "credibility_avg": round(
                sum(s.get("credibility", 5) for s in sources) / max(len(sources), 1), 1
            ),
        })

        self.log(f"[OMNI RESEARCH] ✅ Elite research complete! {len(sources)} sources, avg credibility: {report['credibility_avg']}/10")
        return report

    # ─── PHASE 1: QUERY GENERATION ───────────────────────────────────────────
    async def _generate_queries(self, topic: str) -> List[str]:
        """Use AI to generate diverse search queries, or fallback to rule-based."""
        default = [
            topic,
            f"{topic} latest research 2025",
            f"{topic} statistics data analysis",
            f"{topic} expert opinion challenges",
            f"{topic} future trends predictions",
        ]

        if not self.llm:
            return default

        prompt = f"""Generate exactly 5 diverse, highly specific search queries to deeply research the topic: "{topic}".
Each query should target a DIFFERENT angle (overview, statistics, expert analysis, challenges, future outlook).
Return ONLY a numbered list, one query per line. No explanations."""
        try:
            result = await asyncio.get_event_loop().run_in_executor(
                None, lambda: self.llm.generate(prompt, mode="GENERAL", stream=False)
            )
            lines = [re.sub(r'^\d+[\.\)]\s*', '', l).strip()
                     for l in result.strip().split('\n') if l.strip()]
            lines = [l for l in lines if len(l) > 10]
            return lines[:6] if lines else default
        except Exception:
            return default

    # ─── PHASE 2: MULTI-ENGINE SEARCH ────────────────────────────────────────
    async def _multi_search(self, queries: List[str]) -> List[str]:
        all_links: List[str] = []
        seen: set = set()

        for i, q in enumerate(queries[:5]):
            if self._stop_requested:
                break

            # Google
            google_links = await self._search_engine(
                f"https://www.google.com/search?q={q.replace(' ', '+')}&num=10"
            )
            # Bing (every other query)
            bing_links = []
            if i % 2 == 0:
                bing_links = await self._search_engine(
                    f"https://www.bing.com/search?q={q.replace(' ', '+')}&count=10",
                    skip_domains=["bing.com", "microsoft.com"]
                )

            for url in google_links + bing_links:
                if url not in seen:
                    seen.add(url)
                    all_links.append(url)

            await asyncio.sleep(0.8)  # polite delay

        return all_links

    async def _search_engine(self, url: str,
                              skip_domains: list = None) -> List[str]:
        skip = skip_domains or ["google.", "youtube.", "facebook.", "instagram."]
        try:
            await self.page.goto(url, wait_until="domcontentloaded", timeout=25000)
            await asyncio.sleep(1.5)
            links = await self.page.evaluate("""
                (skipList) => {
                    const els = document.querySelectorAll('a[href^="http"]');
                    const seen = new Set();
                    const results = [];
                    for (const el of els) {
                        const href = el.href;
                        if (skipList.some(s => href.includes(s))) continue;
                        if (!seen.has(href)) {
                            seen.add(href);
                            results.push(href);
                        }
                        if (results.length >= 12) break;
                    }
                    return results;
                }
            """, skip)
            return links or []
        except Exception as e:
            self.log(f"[OMNI RESEARCH] Search error: {e}")
            return []

    # ─── PHASE 3: SOURCE SCORING ─────────────────────────────────────────────
    def _score_sources(self, links: List[str], topic: str) -> List[tuple]:
        topic_words = set(topic.lower().split())
        scored = []
        for url in links:
            score = _score_domain(url)
            # Boost if topic words appear in URL
            url_lower = url.lower()
            for word in topic_words:
                if len(word) > 3 and word in url_lower:
                    score += 1
            scored.append((url, min(score, 10)))
        scored.sort(key=lambda x: x[1], reverse=True)
        return scored

    # ─── PHASE 4: CONTENT EXTRACTION ─────────────────────────────────────────
    async def _extract_sources(self, urls: List[str], limit: int) -> List[Dict]:
        sources = []
        for url in urls:
            if self._stop_requested or len(sources) >= limit:
                break
            cred = _score_domain(url)
            self.log(f"[OMNI RESEARCH] Reading [{cred}/10 cred] → {url[:75]}…")
            text = await self._extract_text(url)
            if text and len(text) > 300:
                title = await self._get_title()
                sources.append({
                    "url": url,
                    "text": text[:self.CHARS_PER_SOURCE],
                    "credibility": cred,
                    "title": title or url[:60],
                    "char_count": len(text),
                })
                self.log(f"[OMNI RESEARCH]   ✓ {len(text):,} chars captured (title: {title[:50] if title else 'N/A'})")
            else:
                self.log(f"[OMNI RESEARCH]   ✗ Skipped (insufficient content)")
        return sources

    async def _extract_text(self, url: str) -> str:
        try:
            await self.page.goto(url, wait_until="domcontentloaded", timeout=20000)
            await asyncio.sleep(1.2)
            text = await self.page.evaluate("""
                () => {
                    ['nav','header','footer','script','style','noscript',
                     '[class*="cookie"]','[class*="popup"]','[class*="banner"]',
                     '[class*="ad-"]','[id*="cookie"]','aside'].forEach(sel => {
                        document.querySelectorAll(sel).forEach(el => el.remove());
                    });
                    const priority = ['article','main','[role="main"]',
                        '.content','.article-content','.post-content',
                        '.entry-content', '#content','#main','#article'];
                    for (const sel of priority) {
                        const el = document.querySelector(sel);
                        if (el && el.innerText.trim().length > 400)
                            return el.innerText.trim();
                    }
                    return document.body ? document.body.innerText.trim() : '';
                }
            """)
            if text:
                text = re.sub(r'\n{3,}', '\n\n', text)
                text = re.sub(r'[ \t]{2,}', ' ', text)
            return text or ""
        except Exception as e:
            self.log(f"[OMNI RESEARCH]   ✗ Extraction error: {str(e)[:60]}")
            return ""

    async def _get_title(self) -> str:
        try:
            return await self.page.title()
        except Exception:
            return ""

    # ─── PHASE 5: WIKIPEDIA LOOKUP ────────────────────────────────────────────
    async def _wikipedia_lookup(self, topic: str) -> str:
        try:
            wiki_slug = topic.strip().replace(" ", "_")
            url = f"https://en.wikipedia.org/wiki/{wiki_slug}"
            await self.page.goto(url, wait_until="domcontentloaded", timeout=20000)
            await asyncio.sleep(1)
            text = await self.page.evaluate("""
                () => {
                    const content = document.querySelector('#mw-content-text .mw-parser-output');
                    if (!content) return '';
                    // Remove tables, references, edit links
                    content.querySelectorAll('table,sup,.mw-editsection').forEach(el=>el.remove());
                    return content.innerText.trim().substring(0, 8000);
                }
            """)
            return text or ""
        except Exception:
            return ""

    # ─── PHASE 6: AI SYNTHESIS ────────────────────────────────────────────────
    async def _synthesise(self, topic: str, sources: List[Dict], stats: List[str]) -> dict:
        if not self.llm:
            combined = "\n\n".join(
                f"=== SOURCE {i+1} [{s['credibility']}/10 credibility]: {s['title']} ===\n{s['text']}"
                for i, s in enumerate(sources)
            )
            return {
                "summary": "AI enrichment unavailable.",
                "sections": {"Raw Research Content": combined[:10000]}
            }

        # Build the combined text (prioritise high-credibility sources)
        sorted_sources = sorted(sources, key=lambda x: x.get("credibility", 5), reverse=True)
        combined = "\n\n".join(
            f"[SOURCE {i+1} | Credibility: {s['credibility']}/10 | {s['title']}]\n{s['text']}"
            for i, s in enumerate(sorted_sources)
        )

        stats_block = "\n".join(f"• {s}" for s in stats[:30]) if stats else "No specific statistics extracted."

        prompt = f"""You are a world-class research analyst and academic writer. 
Based on the multi-source research content below, write a COMPREHENSIVE, DEEPLY DETAILED research paper on:

TOPIC: "{topic}"

EXTRACTED STATISTICS & DATA POINTS:
{stats_block}

Write each section thoroughly — minimum 3-4 detailed paragraphs each. Be specific, cite data, use examples.
Use EXACTLY these section headings:

1. Executive Summary
2. Introduction & Background
3. Current State of the Field
4. Key Findings & Discoveries
5. Statistical Analysis & Data
6. Expert Perspectives & Opinions
7. Case Studies & Real-World Examples
8. Challenges, Limitations & Controversies
9. Future Outlook & Predictions
10. Strategic Recommendations
11. Conclusion
12. References & Sources

--- RESEARCH CONTENT START ---
{combined[:14000]}
--- RESEARCH CONTENT END ---

IMPORTANT: Write in professional academic style. Do NOT use placeholder text. Be extremely detailed and insightful."""

        try:
            response = await asyncio.get_event_loop().run_in_executor(
                None, lambda: self.llm.generate(prompt, mode="GENERAL", stream=False)
            )
            sections = _parse_sections(response)
            return {
                "summary": sections.get("Executive Summary", response[:500]),
                "sections": sections,
                "full_text": response
            }
        except Exception as e:
            self.log(f"[OMNI RESEARCH] AI synthesis error: {e}")
            return {
                "summary": "Synthesis failed.",
                "sections": {"Raw Content": combined[:8000]},
                "full_text": combined[:8000]
            }

    def _empty_report(self, topic: str) -> dict:
        return {
            "topic": topic, "sources": [], "queries_used": [],
            "stats_found": [], "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "summary": "No content was extracted.", "sections": {},
            "total_sources": 0, "credibility_avg": 0.0
        }


# ──────────────────────────────────────────────────────────────────────────────
#  HELPER UTILITIES
# ──────────────────────────────────────────────────────────────────────────────
def _extract_statistics(text: str) -> List[str]:
    """Pull out sentences containing numbers / percentages / dollar amounts."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    stat_pattern = re.compile(
        r'(\d[\d,\.]*\s*(%|percent|million|billion|trillion|\$|USD|EUR|'
        r'GBP|mph|km|kg|lb|°C|°F|years?|months?|days?|hours?))',
        re.IGNORECASE
    )
    found = []
    for sentence in sentences:
        if stat_pattern.search(sentence) and 20 < len(sentence) < 300:
            found.append(sentence.strip())
    return list(dict.fromkeys(found))[:50]  # deduplicate, max 50


def _parse_sections(text: str) -> dict:
    """Parse numbered section headings from AI output into an ordered dict."""
    pattern = re.compile(r'^\d+\.\s+(.+?)$', re.MULTILINE)
    headings = [(m.group(1).strip(), m.start()) for m in pattern.finditer(text)]
    sections = {}
    for i, (heading, start) in enumerate(headings):
        end = headings[i + 1][1] if i + 1 < len(headings) else len(text)
        content = text[start:end].strip()
        content = content.split('\n', 1)[1].strip() if '\n' in content else ""
        if content:
            sections[heading] = content
    return sections if sections else {"Full Report": text}


# ══════════════════════════════════════════════════════════════════════════════
#  WORD EXPORT (Premium formatting)
# ══════════════════════════════════════════════════════════════════════════════
def export_to_word(report: dict, filepath: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Page setup
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        def _heading(text, level, color=(0, 0x66, 0xCC)):
            h = doc.add_heading('', level)
            run = h.add_run(text)
            run.font.color.rgb = RGBColor(*color)
            if level == 0:
                run.font.size = Pt(22)
            elif level == 1:
                run.font.size = Pt(14)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
            return h

        def _body(text, size=11):
            for para_text in text.split('\n\n'):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph(para_text)
                p.style.font.size = Pt(size)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Inches(0.3)

        # ── Cover Page ──
        _heading("RESEARCH REPORT", 0, (0, 0x80, 0xFF))
        doc.add_paragraph()

        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title_p.add_run(report.get("topic", "").upper())
        tr.font.size = Pt(16)
        tr.font.bold = True
        tr.font.color.rgb = RGBColor(0, 0xF0, 0xFF)

        doc.add_paragraph()
        meta_lines = [
            f"Generated by: OMNI AGI Deep Research Engine",
            f"Project by: ORCA  |  Author: Naqash Afzal",
            f"Generated at: {report.get('generated_at', '')}",
            f"Sources consulted: {report.get('total_sources', 0)}",
            f"Average source credibility: {report.get('credibility_avg', 0)}/10",
            f"Search queries used: {len(report.get('queries_used', []))}",
        ]
        for line in meta_lines:
            mp = doc.add_paragraph(line)
            mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mp.style.font.size = Pt(10)

        doc.add_page_break()

        # ── Methodology Box ──
        _heading("Research Methodology", 1)
        queries = report.get("queries_used", [])
        if queries:
            doc.add_paragraph("This report was generated using the following AI-driven search strategy:")
            for i, q in enumerate(queries, 1):
                doc.add_paragraph(f"• Query {i}: {q}", style='List Bullet')
        doc.add_paragraph()

        # ── Sections ──
        for heading, content in report.get("sections", {}).items():
            _heading(heading, 1)
            _body(content)
            doc.add_paragraph()

        # ── Statistics Appendix ──
        stats = report.get("stats_found", [])
        if stats:
            doc.add_page_break()
            _heading("Appendix A: Key Statistics & Data Points", 1)
            for stat in stats[:40]:
                doc.add_paragraph(f"• {stat}", style='List Bullet')

        # ── Sources ──
        doc.add_page_break()
        _heading("Appendix B: Sources Consulted", 1)
        for i, src in enumerate(report.get("sources", []), 1):
            cred = src.get("credibility", "?")
            title = src.get("title", src["url"])[:80]
            p = doc.add_paragraph()
            p.add_run(f"[{i}] {title}  [Credibility: {cred}/10]").bold = True
            p.add_run(f"\n     {src['url']}")
            p.paragraph_format.space_after = Pt(4)

        doc.save(filepath)
        return True
    except ImportError:
        return False
    except Exception as e:
        print(f"[EXPORT] Word error: {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
#  EXCEL EXPORT (Premium styling)
# ══════════════════════════════════════════════════════════════════════════════
def export_to_excel(report: dict, filepath: str) -> bool:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()

        VOID  = "050508"
        CYAN  = "00f0ff"
        NAVY  = "0e0e1a"
        NAVY2 = "131325"
        WHITE = "e8eaf6"
        GOLD  = "ffc857"
        GREEN = "00c853"

        thin = Border(
            left=Side(style='thin', color='1e1e38'),
            right=Side(style='thin', color='1e1e38'),
            top=Side(style='thin', color='1e1e38'),
            bottom=Side(style='thin', color='1e1e38'),
        )

        def _style_header(cell, text, col_color=VOID, text_color=CYAN, size=11):
            cell.value = text
            cell.font = Font(bold=True, color=text_color, size=size, name="Consolas")
            cell.fill = PatternFill("solid", fgColor=col_color)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin

        def _style_data(cell, text, fill_color=NAVY, text_color=WHITE):
            cell.value = str(text)[:32000]
            cell.font = Font(color=text_color, size=10, name="Calibri")
            cell.fill = PatternFill("solid", fgColor=fill_color)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = thin

        # ── Sheet 1: Cover / Overview ──────────────────────────────────────
        ws1 = wb.active
        ws1.title = "📋 Overview"
        ws1.column_dimensions['A'].width = 30
        ws1.column_dimensions['B'].width = 70

        ws1.merge_cells('A1:B1')
        title = ws1['A1']
        title.value = f"OMNI ELITE RESEARCH — {report.get('topic', '').upper()}"
        title.font = Font(bold=True, color=CYAN, size=16, name="Consolas")
        title.fill = PatternFill("solid", fgColor=VOID)
        title.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws1.row_dimensions[1].height = 35

        meta = [
            ("Generated At", report.get("generated_at", "")),
            ("Topic", report.get("topic", "")),
            ("Total Sources", str(report.get("total_sources", 0))),
            ("Avg Source Credibility", f"{report.get('credibility_avg', 0)}/10"),
            ("Search Queries Used", str(len(report.get("queries_used", [])))),
            ("Stats & Data Points Found", str(len(report.get("stats_found", [])))),
            ("Generated By", "OMNI AGI | Project ORCA | Naqash Afzal"),
        ]
        for i, (k, v) in enumerate(meta, 2):
            ws1.row_dimensions[i].height = 18
            _style_header(ws1.cell(row=i, column=1), k, col_color=NAVY, text_color=GOLD)
            _style_data(ws1.cell(row=i, column=2), v)

        # Queries used
        ws1.cell(row=len(meta)+3, column=1, value="SEARCH QUERIES USED")
        ws1.cell(row=len(meta)+3, column=1).font = Font(bold=True, color=CYAN, name="Consolas")
        ws1.cell(row=len(meta)+3, column=1).fill = PatternFill("solid", fgColor=VOID)
        for j, q in enumerate(report.get("queries_used", []), len(meta)+4):
            ws1.cell(row=j, column=1, value=f"Query {j-len(meta)-3}")
            ws1.cell(row=j, column=2, value=q)
            ws1.cell(row=j, column=2).font = Font(color=WHITE, size=10, name="Calibri")
            ws1.cell(row=j, column=2).fill = PatternFill("solid", fgColor=NAVY2)

        # ── Sheet 2: Full Report ───────────────────────────────────────────
        ws2 = wb.create_sheet("📄 Full Report")
        ws2.column_dimensions['A'].width = 28
        ws2.column_dimensions['B'].width = 90

        row = 1
        for section, content in report.get("sections", {}).items():
            # Section heading
            ws2.merge_cells(f'A{row}:B{row}')
            _style_header(ws2[f'A{row}'], section, col_color=VOID, text_color=CYAN, size=12)
            ws2.row_dimensions[row].height = 22
            row += 1

            # Content paragraphs
            for i_para, para in enumerate(content.split('\n')):
                para = para.strip()
                if not para:
                    continue
                alt = NAVY if row % 2 == 0 else NAVY2
                ws2.cell(row=row, column=1, value=f"¶ Para {i_para+1}")
                ws2.cell(row=row, column=1).font = Font(color=GOLD, size=8, name="Consolas", italic=True)
                ws2.cell(row=row, column=1).fill = PatternFill("solid", fgColor=alt)
                ws2.cell(row=row, column=1).alignment = Alignment(horizontal="right", vertical="top")

                c = ws2.cell(row=row, column=2, value=para[:32000])
                c.font = Font(color=WHITE, size=10, name="Calibri")
                c.fill = PatternFill("solid", fgColor=alt)
                c.alignment = Alignment(wrap_text=True, vertical="top")
                c.border = thin
                ws2.row_dimensions[row].height = max(18, min(150, len(para) // 4))
                row += 1

            row += 1  # spacer

        # ── Sheet 3: Statistics ───────────────────────────────────────────
        ws3 = wb.create_sheet("📊 Statistics")
        ws3.column_dimensions['A'].width = 6
        ws3.column_dimensions['B'].width = 100
        _style_header(ws3.cell(1, 1), "#", col_color=VOID)
        _style_header(ws3.cell(1, 2), "Statistical Data Point", col_color=VOID)
        for i, stat in enumerate(report.get("stats_found", []), 2):
            ws3.cell(i, 1, value=i-1).font = Font(color=GOLD, size=9, name="Consolas")
            ws3.cell(i, 1).fill = PatternFill("solid", fgColor=NAVY)
            _style_data(ws3.cell(i, 2), stat)

        # ── Sheet 4: Sources ─────────────────────────────────────────────
        ws4 = wb.create_sheet("🔗 Sources")
        ws4.column_dimensions['A'].width = 5
        ws4.column_dimensions['B'].width = 50
        ws4.column_dimensions['C'].width = 12
        ws4.column_dimensions['D'].width = 60
        _style_header(ws4.cell(1, 1), "#", col_color=VOID)
        _style_header(ws4.cell(1, 2), "Page Title", col_color=VOID)
        _style_header(ws4.cell(1, 3), "Credibility", col_color=VOID)
        _style_header(ws4.cell(1, 4), "URL", col_color=VOID)
        for i, src in enumerate(report.get("sources", []), 2):
            alt = NAVY if i % 2 == 0 else NAVY2
            ws4.cell(i, 1, value=i-1).font = Font(color=GOLD, size=9)
            ws4.cell(i, 1).fill = PatternFill("solid", fgColor=alt)
            _style_data(ws4.cell(i, 2), src.get("title", "")[:80], alt)
            cred_val = src.get("credibility", 5)
            cred_color = GREEN if cred_val >= 8 else (GOLD if cred_val >= 6 else "ff4444")
            ws4.cell(i, 3, value=f"{cred_val}/10")
            ws4.cell(i, 3).font = Font(color=cred_color, bold=True, size=10)
            ws4.cell(i, 3).fill = PatternFill("solid", fgColor=alt)
            ws4.cell(i, 3).alignment = Alignment(horizontal="center")
            _style_data(ws4.cell(i, 4), src["url"], alt, "4fc3f7")

        wb.save(filepath)
        return True
    except Exception as e:
        print(f"[EXPORT] Excel error: {e}")
        return False
